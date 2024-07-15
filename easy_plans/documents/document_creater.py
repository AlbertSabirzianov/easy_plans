import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph

from plans.models import Plan


class DocumentMaker:

    def __init__(self, plan: Plan):
        self.document = Document()
        self.plan = plan

    @property
    def school(self):
        return self.plan.student.work_place.school

    @property
    def user(self):
        return self.plan.student.work_place.user

    @property
    def student(self):
        return self.plan.student

    @property
    def __student_name(self):
        if not self.student.sur_name:
            return self.student.full_name
        return self.student.full_name + ' ' + self.student.sur_name

    @staticmethod
    def make_pr_on_side(pr: Paragraph, side):
        pr.paragraph_format.alignment = side

    def add_empty_paragraphs(self, number: int):
        for _ in range(number):
            self.document.add_paragraph()

    def set_title(self):
        school_pr = self.document.add_heading(
            self.school.name
        )
        self.make_pr_on_side(school_pr, WD_ALIGN_PARAGRAPH.RIGHT)
        self.add_empty_paragraphs(2)

        head = self.document.add_paragraph()
        run = head.add_run("Индивидуальный План")
        run.bold = True
        self.make_pr_on_side(head, WD_ALIGN_PARAGRAPH.CENTER)

        name_pr = self.document.add_paragraph("Ученика (цы) ")
        run = name_pr.add_run(self.student.full_name)
        run.underline = True
        self.make_pr_on_side(
            name_pr,
            WD_ALIGN_PARAGRAPH.CENTER
        )

        dep_pr = self.document.add_paragraph("Отделение ")
        run = dep_pr.add_run(self.plan.department)
        run.underline = True
        self.make_pr_on_side(
            dep_pr,
            WD_ALIGN_PARAGRAPH.CENTER
        )
        self.document.add_page_break()

    def add_header(self, text: str):
        head = self.document.add_paragraph()
        run = head.add_run(text)
        run.bold = True
        self.make_pr_on_side(head, WD_ALIGN_PARAGRAPH.CENTER)

    def add_line(self, text, line):
        pass

    def set_student_page(self):
        self.add_header("Сведения об Учащемся")

        pr = self.document.add_paragraph("Фамилия, Имя, Отчество ")
        run = pr.add_run(self.__student_name)
        run.underline = True

    def make_document(self):
        self.set_title()
        for section in self.document.sections:
            section.page_height = Cm(21)
            section.page_width = Cm(14.8)

    def get_document_stream(self) -> io.BytesIO:
        self.make_document()
        stream = io.BytesIO()
        self.document.save(stream)
        return stream
